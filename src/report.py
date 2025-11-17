import copy
from datetime import datetime
from webbrowser import open_new

import pandas as pd
import streamlit as st
from dateutil.relativedelta import relativedelta
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListItem, ListFlowable
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.pdfmetrics import registerFontFamily
from reportlab.pdfgen import canvas
from io import BytesIO
from ContentBlock import ContentBlock, Alignment
from docx.oxml.ns import qn
import base64
import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# Refer to ContentBlock.py for list of useful functions. Use a list within the report_content to activate a bullet list
# {DateRange} gets replaced with the date range
# {PrevMonth} gets replaced with the month before the beginning month
# {TwoPrevMonth} gets replaced with two months before the beginning month

def get_satisfaction_percentages(sheet):
    df = pd.read_excel('data/rawdata/CustomerSurveyResponses.xlsx', sheet_name=sheet)
    df['answerValues'] = pd.to_numeric(df['answerValues'], errors='coerce')
    satisfaction_df = df[df['questionLabel'].str.contains('satisfaction - overall', case=False, na=False)]
    satisfied = satisfaction_df[satisfaction_df['answerValues'] >= 4]['respondentId'].nunique()
    total = satisfaction_df['respondentId'].nunique()
    percentage_satisfied = (satisfied / total) * 100 if total > 0 else 0
    return percentage_satisfied, total

def get_answer_percentages(sheet, question_label, target_answers):
    df = pd.read_excel("data/rawdata/CustomerSurveyResponses.xlsx", sheet_name=sheet)
    question_df = df[df["questionLabel"] == question_label].copy()
    question_df["answerLabels"] = question_df["answerLabels"].astype(str)
    if isinstance(target_answers, str):
        target_answers = [target_answers]
    respondent_groups = question_df.groupby("respondentId")
    total_respondents = respondent_groups.ngroups

    def respondent_has_target(group):
        return any(
            group["answerLabels"].str.contains(answer, case=False, na=False).any()
            for answer in target_answers
        )

    respondents_with_answer = sum(respondent_groups.apply(respondent_has_target))
    percentage = (respondents_with_answer / total_respondents) * 100 if total_respondents > 0 else 0
    return respondents_with_answer, total_respondents, percentage

def get_report_content():
    # Load the details file, which has the most details of any advertising email dataset and is mandatory in all data
    try:
        details_df = pd.read_csv("data/convertedcsv/Advertising_Email_Engagement/Advertising_Email_Engagement.xlsx-Email_Engagement_Details.csv")

        # Ensure the file isn’t empty
        if details_df.empty or details_df.shape[0] < 2:
            return [ContentBlock("No data found. Please enter all content manually or provide non-empty data sheets.")]
    except (FileNotFoundError, pd.errors.EmptyDataError):
        return [ContentBlock("Advertising Email Engagement Details data file is not found. Please enter all content manually or provide the data file.")]

    # If there is at least one email in the database, then proceed with the rest of the function

    # GET CURRENT MONTH OPEN RATE
    open_rate = None
    try:
        open_rate_df = pd.read_csv("data/convertedcsv/Advertising_Email_Engagement/Advertising_Email_Engagement.xlsx-Open_Rate.csv")
        # Check if column exists and has at least one non-NaN value
        if "Open Rate" in open_rate_df.columns and not open_rate_df["Open Rate"].isna().all():
            open_rate = open_rate_df["Open Rate"].iloc[0]
            open_rate_str = f"{open_rate:.2%}"
        else:
            open_rate_str = "[Enter Open Rate]"
    except (FileNotFoundError, pd.errors.EmptyDataError):
        open_rate_str = "[Enter Open Rate]"

    # GET CLICK THROUGH RATE
    try:
        click_to_open_rate_df = pd.read_csv("data/convertedcsv/Advertising_Email_Engagement/Advertising_Email_Engagement.xlsx-Click_to_Open_Rate.csv")
        # Check if column exists and has at least one non-NaN value
        if "Click To Open Rate" in click_to_open_rate_df.columns and not click_to_open_rate_df["Click To Open Rate"].isna().all():
            click_to_open_rate = click_to_open_rate_df["Click To Open Rate"].iloc[0]
            if click_to_open_rate is not None and open_rate is not None:
                click_through_rate = click_to_open_rate * open_rate
                click_through_rate_str = f"{click_through_rate:.2%}"
            else:
                click_through_rate_str = "[Enter Click Through Rate]"
        else:
            click_through_rate_str = "[Enter Click Through Rate]"
    except (FileNotFoundError, pd.errors.EmptyDataError):
        click_through_rate_str = "[Enter Click Through Rate]"

    # GET HIGHLIGHTS
    click_sorted = details_df.sort_values(by="Click Rate", ascending=False).reset_index(drop=True)
    open_sorted = details_df.sort_values(by="Open Rate", ascending=False).reset_index(drop=True)

    highlights = []

    top_click = click_sorted.iloc[0]
    top_open = open_sorted.iloc[0]

    if top_click["Message Name"] == top_open["Message Name"]:
        highlights.append(top_click)
        if len(click_sorted) > 2:
            second_click = click_sorted.iloc[1]
            second_open = open_sorted.iloc[1]
            if second_click["Message Name"] == second_open["Message Name"]:
                highlights.append(second_click)
            else:
                highlights.append(second_open)
                highlights.append(second_click)
    else:
        highlights.append(top_click)
        highlights.append(top_open)

    highlight_content_blocks = []
    for highlight in highlights:
        highlight_content_block = ContentBlock(f"\"{highlight['Message Name']}\" {highlight['Open Rate']:.2%} OPR {highlight['Click Rate']:.2%} CTR", indent=True, new_line=True)
        highlight_content_blocks.append(highlight_content_block)

    # GET CAMPAIGNS
    # Drop rows where 'Campaign' is NaN or empty
    campaigns_df = details_df.dropna(subset=["Campaign"])
    campaigns_df = campaigns_df[campaigns_df["Campaign"].astype(str).str.strip() != ""]

    campaign_content_blocks = []
    if campaigns_df.empty:
        campaign_content_blocks.append(ContentBlock("No campaign data for this time period.", bold=True, color="#000071", new_line=True))
    else:
        # Convert dates
        campaigns_df["Send Date"] = pd.to_datetime(campaigns_df["Send Date"], errors="coerce")

        # Drop rows with invalid dates
        campaigns_df = campaigns_df.dropna(subset=["Send Date"])

        # Group by Campaign
        for campaign, group in campaigns_df.groupby("Campaign"):
            start_date = group["Send Date"].min().strftime("%d-%b")
            end_date = group["Send Date"].max().strftime("%d-%b")
            avg_open = group["Open Rate"].mean(skipna=True)
            avg_click = group["Click Rate"].mean(skipna=True)

            campaign_content_blocks.append(ContentBlock(f"{campaign.title()}", bold=True, color="#000071"))
            campaign_content_blocks.append(ContentBlock(f" ({start_date} - {end_date} TY):", new_line=True))
            campaign_content_blocks.append(ContentBlock(f"Average OPR: {avg_open:.1%} I Average CTR: {avg_click:.1%}", indent=True, new_line=True))

    # GET MAIN EXCHANGE SURVEYS
    main_stores_percentage_satisfied, main_stores_total = get_satisfaction_percentages("Main Stores")
    _, _, main_stores_percentage_advertised = get_answer_percentages("Main Stores", "MCX_Primary Reason", "Shopping sales that were advertised")
    _, _, main_stores_percentage_picking_up_needed_supplies = get_answer_percentages("Main Stores", "MCX_Primary Reason", "Picking up needed supplies")
    marine_marts_percentage_satisfied, marine_marts_total = get_satisfaction_percentages("Marine Marts")
    marine_marts_didnt_find_count, _, marine_marts_percentage_didnt_find = get_answer_percentages("Marine Marts", "Items Found", ["Partially", "No"])

    # PUT TOGETHER REPORT
    email_findings = [
         ContentBlock("The industry standard open rate (OPR) for emails in retail is 15-25% - MCX average for {DateRange} was "),
         ContentBlock(f"{open_rate_str}", bold=True, color="#000071"),
         ContentBlock(" ([Enter Last Month Open Rate] in {PrevMonth}, [Enter Open Rate From Two Months Ago] in {TwoPrevMonth})", new_line=True),

         ContentBlock("The industry standard click through rate (CTR) is 1-5% - MCX average for {DateRange} was "),
         ContentBlock(f"{click_through_rate_str}", bold=True, color="#000071"),
         ContentBlock(" ([Enter Last Month CTR Rate] in {PrevMonth}, [Enter CTR From Two Months Ago] in {TwoPrevMonth})", new_line=True),

         ContentBlock("Performance by email blast (highlights):", new_line=True),
    ]

    for highlight_content_block in highlight_content_blocks:
        email_findings.append(highlight_content_block)
    for campaign_content_block in campaign_content_blocks:
        email_findings.append(campaign_content_block)

    report_content = [
        ContentBlock("{DateRange} MCCS Marketing Analytics Assessment",
                     bold=True,
                     font_size=15,
                     color="#24446C",
                     alignment=Alignment.CENTER,
                     new_paragraph=True),
        ContentBlock("Purpose:",
                     bold=True,
                     underline=True),
        ContentBlock(" To support leaders and subject matter experts in their decision-making process while aiming to generate smart revenue, increase brand affinity, and reduce stress on Marines and their families.",
                     new_paragraph=True),
        ContentBlock("Findings - Review of digital performance, advertising campaigns, and sales:",
                     bold=True,
                     new_line=True),
        email_findings,
        ContentBlock("Findings - Review of Main Exchanges, Marine Marts, and MCHS CSAT Surveys and Google Reviews:",
                     bold=True,
                     new_line=True),
        [
            ContentBlock(f"{round(main_stores_percentage_satisfied)}% of {main_stores_total} Main Exchange survey respondents reported overall satisfaction with their experience.", color="#000071", bold=True, new_line=True),
            ContentBlock(f"{main_stores_percentage_advertised:.1f}% said they were shopping sales that were advertised. {main_stores_percentage_picking_up_needed_supplies:.1f}% were picking up needed supplies.", indent=True, new_line=True),
            ContentBlock(f"{round(marine_marts_percentage_satisfied)}% of {marine_marts_total} Marine Mart survey respondents reported overall satisfaction with their experience.", color="#000071", bold=True, new_line=True),
            ContentBlock(f"{marine_marts_didnt_find_count} out of {marine_marts_total} ({marine_marts_percentage_didnt_find:.1f}%) customers were unable to purchase everything they intended.", indent=True, new_line=True),
            ContentBlock("There were 392 MCHS survey respondents, with an average CSAT score of 91.8.", color="#000071", bold=True, new_line=True),
            ContentBlock("In September 2023, there were 603 survey respondents with an average CSAT score of 88.1", indent=True, new_line=True),
            ContentBlock("40.8% of respondents were T AD/TDY. 32.1% of respondents were traveling for leisure.", indent=True, new_line=True),
            ContentBlock("Respondents traveling for leisure averaged a CSAT score of 90.1. Respondents T AD/TDY averaged 87.6.", indent=True, new_line=True),
            ContentBlock("All time ", indent=True),
            ContentBlock("Google Reviews", bold=True, color="#000071"),
            ContentBlock(" have an average rating of 4.4 out of 5 and there has been a total of 10,637 reviews.", new_paragraph=True),
        ],
        ContentBlock("Assessments",
                     bold=True,
                     new_line=True),
    ]
    return report_content

def add_raw_output_to_report_content(raw_assessments, raw_social_media_report):
    # Process assessments and social_media_report raw output
    assessments = raw_assessments.split("\n")
    cleaned_assessments = []
    for assessment in assessments:
        cleaned_assessment = assessment.partition(".")[2].strip()
        if len(cleaned_assessment.strip()) > 0:
            cleaned_assessments.append(cleaned_assessment)

    social_media_findings = raw_social_media_report.split("\n\n")
    cleaned_social_media_findings = []
    for finding in social_media_findings:
        if len(finding.strip()) > 0:
            cleaned_finding = finding.partition(".")[2].strip()
            parts = cleaned_finding.split("\n")
            if len(parts) == 2:
                cleaned_line = parts[0] + " (Evidence: " + parts[1].strip()[len("- Evidence:"):].strip() + ")"
                cleaned_social_media_findings.append(cleaned_line)
            else:
                cleaned_social_media_findings.append(cleaned_finding)

    # Turn it into content blocks
    assessments_content = []
    for i, assessment in enumerate(cleaned_assessments):
        if i == len(cleaned_assessments) - 1:
            content_block = ContentBlock(assessment, new_paragraph=True)
        else:
            content_block = ContentBlock(assessment, new_line=True)
        assessments_content.append(content_block)

    social_media_content = []
    for i, finding in enumerate(cleaned_social_media_findings):
        if i == len(cleaned_social_media_findings) - 1:
            content_block = ContentBlock(finding, new_paragraph=True)
        else:
            content_block = ContentBlock(finding, new_line=True)
        social_media_content.append(content_block)

    # Build the report using the existing template + new Ollama Information
    current_report = get_report_content()
    current_report.append(assessments_content)
    current_report.append(ContentBlock("{DateRange} MCX Email Highlights",
                                       bold=True,
                                       font_size=15,
                                       color="#24446C",
                                       underline=True,
                                       alignment=Alignment.CENTER,
                                       new_paragraph=True)),
    current_report.append(ContentBlock("{DateRange} MCX Social Media Highlights",
                                       bold=True,
                                       font_size=15,
                                       color="#24446C",
                                       underline=True,
                                       alignment=Alignment.CENTER,
                                       new_paragraph=True)),
    current_report.append(social_media_content),
    current_report.append(ContentBlock("{DateRange} MCX Customer Satisfaction Highlights",
                                       bold=True,
                                       font_size=15,
                                       color="#24446C",
                                       underline=True,
                                       alignment=Alignment.CENTER,
                                       new_paragraph=True)),
    return current_report

def parse_time(begin, end):
    begin_dt = datetime.strptime(begin, "%b %Y")
    end_dt = datetime.strptime(end, "%b %Y")

    # Build replacement strings
    if begin_dt == end_dt:
        date_range = begin_dt.strftime("%B %Y")
    else:
        date_range = f"{begin_dt.strftime('%B')} - {end_dt.strftime('%B %Y')}"

    prev_month = (begin_dt - relativedelta(months=1)).strftime("%B")
    two_prev_month = (begin_dt - relativedelta(months=2)).strftime("%B")
    return date_range, prev_month, two_prev_month

def replace_placeholders(current_report, beginning_month, ending_month):
    # Get month strings for replacement
    date_range, prev_month, two_prev_month = parse_time(beginning_month, ending_month)

    # Replace strings
    for item in current_report:
        if isinstance(item, ContentBlock):
            item.text = item.text.replace("{DateRange}", date_range)
            item.text = item.text.replace("{PrevMonth}", prev_month)
            item.text = item.text.replace("{TwoPrevMonth}", two_prev_month)
        elif isinstance(item, list):
            for block in item:
                if isinstance(block, ContentBlock):
                    block.text = block.text.replace("{DateRange}", date_range)
                    block.text = block.text.replace("{PrevMonth}", prev_month)
                    block.text = block.text.replace("{TwoPrevMonth}", two_prev_month)

def content_block_to_formatted_text(block):
    tagged = block.text
    if block.bold:
        tagged = f"<b>{tagged}</b>"
    if block.underline:
        tagged = f"<u>{tagged}</u>"
    if block.font_size != 10:
        tagged = f"<font size={block.font_size}>{tagged}</font>"
    if block.color:
        tagged = f"<font color='{block.color}'>{tagged}</font>"
    return tagged

def pdf_bullet_list(styles, bullet_points):
    bullet_items = []
    tagged = ""
    indent_block = False

    for bullet in bullet_points:
        tagged += content_block_to_formatted_text(bullet)
        if bullet.indent:
            indent_block = True

        if bullet.new_paragraph or bullet.new_line:
            paragraph = Paragraph(tagged, styles["Normal"])
            item = ListItem(paragraph, bulletColor='black',
                            leftIndent=36 if indent_block else 18)
            bullet_items.append(item)
            tagged = ""
            indent_block = False

    return ListFlowable(
        bullet_items,
        bulletType="bullet",
        bulletFontName="Montserrat-Medium",
        bulletFontSize=10,
        spaceBefore=12
    )

def create_pdf(beginning_month, ending_month, assessments, social_media_report):
    # Get current report (combines template + Ollama)
    current_report = add_raw_output_to_report_content(assessments, social_media_report)

    # Replace all placeholders
    replace_placeholders(current_report, beginning_month, ending_month)

    # Register fonts and font family
    pdfmetrics.registerFont(TTFont("Montserrat-Light", "src/fonts/Montserrat/static/Montserrat-Light.ttf"))
    pdfmetrics.registerFont(TTFont("Montserrat-Medium", "src/fonts/Montserrat/static/Montserrat-Medium.ttf"))

    registerFontFamily("Montserrat", normal="Montserrat-Light", bold="Montserrat-Medium")

    # Create and set up doc
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=36,  # 0.5 inch on all sides
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    # Styles and fonts
    styles = getSampleStyleSheet()
    styles['Normal'].fontName = "Montserrat-Light"
    center_style = ParagraphStyle(
        "Title",
        fontName="Montserrat-Medium",
        leading=12,
        spaceAfter=10,
        alignment=1
    )
    indent_style = ParagraphStyle(
        "Indent",
        leftIndent=36,
    )

    # Process through current_report list and inject reportlab paragraphs into content_list
    content_list = []
    tagged = ""
    indent_block = False
    for block in current_report:
        if isinstance(block, ContentBlock):
            tagged = tagged + content_block_to_formatted_text(block)
            if block.indent:
                indent_block = True
            if block.alignment == Alignment.CENTER:
                paragraph = Paragraph(tagged, center_style)
                content_list.append(paragraph)
                content_list.append(Spacer(1, 12))
                tagged = ""
                indent_block = False
            elif block.new_line:
                style = indent_style if indent_block else styles["Normal"]
                paragraph = Paragraph(tagged, style)
                content_list.append(paragraph)
                tagged = ""
                indent_block = False
            elif block.new_paragraph:
                style = indent_style if indent_block else styles["Normal"]
                paragraph = Paragraph(tagged, style)
                content_list.append(paragraph)
                content_list.append(Spacer(1, 6))
                tagged = ""
                indent_block = False
        elif isinstance(block, list):
            bullet_list = pdf_bullet_list(styles, block)
            content_list.append(bullet_list)
            content_list.append(Spacer(1, 12))

    # Use content list to build a doc
    doc.build(content_list)

    # Set seeker to the beginning of the doc and return the buffer
    buffer.seek(0)
    return buffer

def create_doc(beginning_month, ending_month, assessments, social_media_report):
    # Get current report (combines template + Ollama)
    current_report = add_raw_output_to_report_content(assessments, social_media_report)

    # Replace all placeholders
    replace_placeholders(current_report, beginning_month, ending_month)

    # Set up document
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    def render_run(run, block):
        run.text = block.text
        run.bold = block.bold
        run.underline = block.underline
        run.font.size = Pt(block.font_size)
        run.font.color.rgb = block.get_rgb_color()
        run.font.name = "Montserrat"
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')  # fallback if Montserrat is not available

    def add_paragraph_from_blocks(blocks, style=None, is_bullet=False, is_indent=False):
        para = document.add_paragraph(style=style)
        if is_bullet:
            para.style = 'List Bullet 2' if is_indent else 'List Bullet'

        para.paragraph_format.line_spacing = Pt(12)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(0)

        first_block = blocks[0]
        if first_block.alignment == Alignment.CENTER:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        for i, block in enumerate(blocks):
            run = para.add_run()
            render_run(run, block)
            if block.new_line and i < len(blocks) - 1:
                run.add_break(WD_BREAK.LINE)

    current_line = []
    indent_line = False
    for item in current_report:
        if isinstance(item, list):
            bullet_item = []
            for block in item:
                bullet_item.append(block)
                if block.indent:
                    indent_line = True
                if block.new_line or block.new_paragraph:
                    if bullet_item:
                        add_paragraph_from_blocks(bullet_item, is_bullet=True, is_indent=indent_line)
                        bullet_item = []
                        indent_line = False
        else:
            current_line.append(item)
            if item.new_line or item.new_paragraph:
                add_paragraph_from_blocks(current_line)
                current_line = []
    if current_line:
        add_paragraph_from_blocks(current_line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer